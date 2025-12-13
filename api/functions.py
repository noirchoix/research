import os
from typing import List, Optional, Tuple, Literal
from pathlib import Path
from pypdf import PdfReader
from docx import Document as DocxDocument
from jinja2 import Template
from dotenv import load_dotenv
import httpx
import shutil
import subprocess
import requests

load_dotenv()

BASE_DIR = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "outputs"
for d in (UPLOAD_DIR, OUTPUT_DIR):
    d.mkdir(parents=True, exist_ok=True)

SERPAPI_API_KEY = os.getenv("SERPAPI_API_KEY", "")

# ============================================================
# UTIL: EXTRACT TEXT FROM PDF/DOCX
# ============================================================

def extract_text_from_pdf(path: Path) -> Tuple[str, str]:
    """Return (title, text)."""
    reader = PdfReader(str(path))
    text_pages = []
    title = ""
    try:
        meta = reader.metadata
        title_candidate = ""
        if meta:
            # metadata may be a dict-like mapping or an object with attributes
            if isinstance(meta, dict):
                # PDF metadata keys can be like "/Title" or "Title"
                title_candidate = meta.get("/Title") or meta.get("Title") or ""
            else:
                title_candidate = getattr(meta, "title", "") or getattr(meta, "Title", "") or ""
        title = (title_candidate or "").strip()
    except Exception:
        title = ""

    for page in reader.pages:
        page_text = page.extract_text() or ""
        text_pages.append(page_text)

    full_text = "\n".join(text_pages)
    if not title:
        # crude heuristic: first non-empty line of first page
        for line in text_pages[0].splitlines():
            if line.strip():
                title = line.strip()[:200]
                break
    if not title:
        title = path.stem
    return title, full_text


def extract_text_from_docx(path: Path) -> Tuple[str, str]:
    doc = DocxDocument(str(path))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paras)
    # title heuristic: first non-empty paragraph
    title = paras[0][:200] if paras else path.stem
    return title, full_text


# ============================================================
# FILE GENERATORS (.txt, .docx, .tex, .pdf)
# ============================================================

def write_txt(text: str, prefix: str) -> Path:
    path = OUTPUT_DIR / f"{prefix}.txt"
    path.write_text(text, encoding="utf-8")
    return path


def write_docx(text: str, prefix: str, title: str) -> Path:
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    for paragraph in text.split("\n\n"):
        doc.add_paragraph(paragraph.strip())
    path = OUTPUT_DIR / f"{prefix}.docx"
    doc.save(str(path))
    return path


# ---------- LaTeX helpers (from your templating/tectonic idea) ----------

def _tectonic_supports_new_cli(tectonic_path: str) -> bool:
    try:
        help_out = subprocess.run(
            [tectonic_path, "-X", "compile", "--help"],
            capture_output=True, text=True, check=False
        ).stdout.lower()
        return "--outdir" in help_out
    except Exception:
        return False


def compile_tex_to_pdf(tex_path: str) -> Tuple[bool, str]:
    tex_file = Path(tex_path)
    workdir = tex_file.parent
    outdir = workdir

    tectonic = shutil.which("tectonic")
    pdflatex = shutil.which("pdflatex")

    try:
        if tectonic:
            if _tectonic_supports_new_cli(tectonic):
                cmd = [
                    tectonic, "-X", "compile",
                    "--outdir", str(outdir),
                    "--keep-intermediates", "--keep-logs",
                    str(tex_file),
                ]
            else:
                cmd = [tectonic, "-o", str(outdir), str(tex_file)]
            proc = subprocess.run(cmd, cwd=str(workdir), capture_output=True, text=True)
        elif pdflatex:
            cmd = [
                pdflatex, "-interaction=nonstopmode", "-halt-on-error",
                "-output-directory", str(outdir),
                str(tex_file),
            ]
            proc = subprocess.run(cmd, cwd=str(workdir), capture_output=True, text=True)
        else:
            return False, "No LaTeX engine found on PATH. Install Tectonic or MiKTeX."

        if proc.returncode != 0:
            return False, (proc.stdout or "") + "\n" + (proc.stderr or "")

        pdf_guess = workdir / (tex_file.stem + ".pdf")
        if pdf_guess.exists():
            return True, str(pdf_guess)

        pdf_alt = Path(str(tex_file).replace(".tex", ".pdf"))
        if pdf_alt.exists():
            return True, str(pdf_alt)

        return False, "Compilation reported success but PDF not found."
    except Exception as e:
        return False, f"Exception during LaTeX compile: {e}"


def write_latex_math_proof(body: str, prefix: str, title: str) -> Path:
    template_str = r"""
\documentclass{article}
\usepackage[utf8]{inputenc}
\usepackage{amsmath, amssymb, amsthm}
\usepackage{geometry}
\geometry{margin=1in}

\begin{document}
\section*{Mathematical Treatment of {{ title }}}

{{ body }}

\end{document}
"""
    t = Template(template_str)
    tex_content = t.render(title=title, body=body)
    tex_path = OUTPUT_DIR / f"{prefix}.tex"
    tex_path.write_text(tex_content, encoding="utf-8")
    return tex_path


def write_latex_research_summary(body: str, prefix: str, title: str) -> Path:
    template_str = r"""
\documentclass{article}
\usepackage[utf8]{inputenc}
\usepackage{geometry}
\geometry{margin=1in}

\begin{document}
\section*{Research Summary: {{ title }}}

{{ body }}

\end{document}
"""
    t = Template(template_str)
    tex_content = t.render(title=title, body=body)
    tex_path = OUTPUT_DIR / f"{prefix}.tex"
    tex_path.write_text(tex_content, encoding="utf-8")
    return tex_path


# ============================================================
# RELATED ARTICLES (SerpAPI / Google Scholar)
# ============================================================

def retrieve_scholar_articles(
    professional_fields: List[str],
    date_window: Tuple[int, int],
    num_of_articles: int = 10,
    language: str = "en",
    include_patents: bool = False,
    api_key: str = "",
    extra_params: Optional[dict] = None
) -> List[dict]:
    if not api_key:
        raise ValueError("You must provide a valid SerpApi api_key.")
    base_url = "https://serpapi.com/search"
    results: List[dict] = []

    for field in professional_fields:
        params = {
            "engine": "google_scholar",
            "q": field,
            "as_ylo": date_window[0],
            "as_yhi": date_window[1],
            "hl": language,
            "as_sdt": "0" if not include_patents else "0,7",
            "num": min(num_of_articles, 20),
            "api_key": api_key,
            "output": "json",
        }
        if extra_params:
            params.update(extra_params)

        resp = requests.get(base_url, params=params)
        data = resp.json()
        if data.get("search_metadata", {}).get("status") != "Success":
            raise RuntimeError(f"Search failed for field '{field}': {data}")

        for item in data.get("organic_results", []):
            pub_info = item.get("publication_info", {}).get("summary", "")
            year = None
            for token in pub_info.split():
                if token.isdigit() and len(token) == 4:
                    year = token
                    break
            results.append({
                "field": field,
                "title": item.get("title"),
                "link": item.get("link"),
                "authors_info": pub_info,
                "publication_year": year,
                "snippet": item.get("snippet"),
            })
    return results[:num_of_articles]
